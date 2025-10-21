"""
financial_report_agent.py

A robust financial report generation agent with proper error handling,
logging, and separation of concerns.
"""

import os
import io
import json
import logging
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, asdict
from datetime import datetime, timezone
from enum import Enum

import numpy as np
import pandas as pd
import yfinance as yf
import matplotlib.pyplot as plt

from langchain.tools import Tool
from langchain.chains import LLMChain
from langchain_openai import ChatOpenAI
from langchain_core.prompts import PromptTemplate
from langchain.agents import initialize_agent, AgentType, Tool as LC_Tool

# PDF generation - install with: pip install markdown reportlab
try:
    import markdown
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PDF libraries not available. Install with: pip install markdown reportlab")

# DOCX generation - install with: pip install python-docx
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("DOCX libraries not available. Install with: pip install python-docx")

# ---------- Logging Configuration ----------
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ---------- Configuration ----------
@dataclass
class Config:
    """Application configuration."""
    openai_api_key: str
    openai_base_url: str = "https://api.openai.com/v1"
    model_name: str = "gpt-4o"
    default_period: str = "1y"
    max_retries: int = 3
    request_timeout: int = 30
    
    @classmethod
    def from_env(cls) -> 'Config':
        """Load configuration from environment variables."""
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("OPENAI_API_KEY environment variable must be set")
        
        return cls(
            openai_api_key=api_key,
            openai_base_url=os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1"),
            model_name=os.getenv("OPENAI_MODEL", "gpt-4o")
        )

# ---------- Data Models ----------
class OutputFormat(Enum):
    """Supported output formats."""
    MARKDOWN = "markdown"
    PDF = "pdf"
    DOCX = "docx"

@dataclass
class ParsedRequest:
    """Structured representation of user request."""
    tickers: List[str]
    period: str
    metrics: List[str]
    output_format: OutputFormat
    extras: Optional[Dict[str, Any]] = None
    
    def validate(self) -> None:
        """Validate parsed request data."""
        if not self.tickers:
            raise ValueError("At least one ticker must be specified")
        if len(self.tickers) > 10:
            raise ValueError("Maximum 10 tickers allowed")
        for ticker in self.tickers:
            if not ticker or not ticker.strip():
                raise ValueError("Invalid ticker: empty string")

@dataclass
class TickerAnalysis:
    """Analysis results for a single ticker."""
    ticker: str
    csv_path: str
    chart_path: str
    latest_close: float
    avg_daily_return: float
    volatility: float
    ratios: Dict[str, Optional[float]]
    sample_data: List[Dict[str, Any]]
    error: Optional[str] = None

@dataclass
class ReportMetadata:
    """Metadata for generated report."""
    final_markdown_path: str
    charts: List[str]
    analyses: Dict[str, TickerAnalysis]
    review_issues: List[str]
    generated_at: str

# ---------- Custom Exceptions ----------
class DataFetchError(Exception):
    """Raised when data fetching fails."""
    pass

class ParsingError(Exception):
    """Raised when request parsing fails."""
    pass

class AnalysisError(Exception):
    """Raised when analysis fails."""
    pass

# ---------- Utility Functions ----------
def normalize_period(user_period: Optional[str]) -> str:
    """
    Convert natural language time periods into valid yfinance periods.
    
    Args:
        user_period: User-specified time period
        
    Returns:
        Valid yfinance period string
    """
    if not user_period:
        return "1y"
    
    p = user_period.strip().lower()
    valid_periods = ["1d", "5d", "1mo", "3mo", "6mo", "1y", "2y", "5y", "10y", "ytd", "max"]
    
    if p in valid_periods:
        return p
    
    # Natural language mappings
    mappings = {
        "quarter": "3mo",
        "month": "1mo",
        "year": "1y",
        "annual": "1y",
        "ytd": "ytd"
    }
    
    for key, value in mappings.items():
        if key in p:
            return value
    
    logger.warning(f"Unknown period '{user_period}', defaulting to 1y")
    return "1y"

def safe_float_conversion(value: Any) -> Optional[float]:
    """Safely convert value to float."""
    try:
        return float(value) if value is not None else None
    except (ValueError, TypeError):
        return None

# ---------- Data Fetching ----------
class DataFetcher:
    """Handles all external data fetching operations."""
    
    def __init__(self, timeout: int = 30):
        self.timeout = timeout
    
    def fetch_price_history(self, ticker: str, period: str = "2y") -> pd.DataFrame:
        """
        Fetch historical price data via yfinance.
        
        Args:
            ticker: Stock ticker symbol
            period: Time period (e.g., '1y', '2y')
            
        Returns:
            DataFrame with historical price data
            
        Raises:
            DataFetchError: If data fetching fails
        """
        if not ticker or not ticker.strip():
            raise ValueError("Ticker cannot be empty")
        
        ticker = ticker.strip().upper()
        logger.info(f"Fetching price history for {ticker} (period: {period})")
        
        try:
            t = yf.Ticker(ticker)
            hist = t.history(period=period, auto_adjust=False)
            
            if hist.empty:
                raise DataFetchError(f"No data returned for {ticker}")
            
            hist = hist.reset_index()
            hist["ticker"] = ticker
            
            logger.info(f"Successfully fetched {len(hist)} rows for {ticker}")
            return hist
            
        except Exception as e:
            logger.error(f"Failed to fetch price history for {ticker}: {e}")
            raise DataFetchError(f"Failed to fetch data for {ticker}: {str(e)}")
    
    def fetch_financials(self, ticker: str) -> Dict[str, Any]:
        """
        Fetch financial statements via yfinance.
        
        Args:
            ticker: Stock ticker symbol
            
        Returns:
            Dictionary with financial statements
            
        Raises:
            DataFetchError: If data fetching fails
        """
        if not ticker or not ticker.strip():
            raise ValueError("Ticker cannot be empty")
        
        ticker = ticker.strip().upper()
        logger.info(f"Fetching financials for {ticker}")
        
        try:
            t = yf.Ticker(ticker)
            
            # Fetch statements with individual error handling
            result = {}
            
            try:
                quarterly_income = t.quarterly_income_stmt
                result["income_stmt"] = quarterly_income.to_dict() if quarterly_income is not None else {}
            except Exception as e:
                logger.warning(f"Failed to fetch income statement for {ticker}: {e}")
                result["income_stmt"] = {}
            
            try:
                quarterly_balance = t.quarterly_balance_sheet
                result["balance_sheet"] = quarterly_balance.to_dict() if quarterly_balance is not None else {}
            except Exception as e:
                logger.warning(f"Failed to fetch balance sheet for {ticker}: {e}")
                result["balance_sheet"] = {}
            
            try:
                quarterly_cash = t.quarterly_cashflow
                result["cashflow"] = quarterly_cash.to_dict() if quarterly_cash is not None else {}
            except Exception as e:
                logger.warning(f"Failed to fetch cashflow for {ticker}: {e}")
                result["cashflow"] = {}
            
            return result
            
        except Exception as e:
            logger.error(f"Failed to fetch financials for {ticker}: {e}")
            raise DataFetchError(f"Failed to fetch financials for {ticker}: {str(e)}")

# ---------- Analysis ----------
class FinancialAnalyzer:
    """Handles financial data analysis."""
    
    @staticmethod
    def compute_metrics(df_prices: pd.DataFrame) -> pd.DataFrame:
        """
        Compute technical and statistical metrics.
        
        Args:
            df_prices: DataFrame with price data
            
        Returns:
            DataFrame with computed metrics
            
        Raises:
            AnalysisError: If analysis fails
        """
        if df_prices.empty:
            raise AnalysisError("Cannot analyze empty DataFrame")
        
        if "Date" not in df_prices.columns or "Close" not in df_prices.columns:
            raise AnalysisError("DataFrame must have 'Date' and 'Close' columns")
        
        try:
            df = df_prices.copy()
            df['Date'] = pd.to_datetime(df['Date'])
            df = df.sort_values('Date')
            
            # Price metrics
            df['close'] = df['Close']
            df['daily_return'] = df['close'].pct_change()
            
            # Moving averages with configurable windows
            df['30d_ma'] = df['close'].rolling(window=30, min_periods=5).mean()
            df['50d_ma'] = df['close'].rolling(window=50, min_periods=10).mean()
            
            # Volatility (20-day rolling standard deviation of returns)
            df['volatility'] = df['daily_return'].rolling(window=20, min_periods=5).std()
            
            logger.info(f"Computed metrics for {len(df)} rows")
            return df
            
        except Exception as e:
            logger.error(f"Failed to compute metrics: {e}")
            raise AnalysisError(f"Analysis failed: {str(e)}")
    
    @staticmethod
    def compute_ratios(financials: Dict[str, Any]) -> Dict[str, Optional[float]]:
        """
        Compute key financial ratios from statements.
        
        Args:
            financials: Dictionary with financial statements
            
        Returns:
            Dictionary of computed ratios
        """
        ratios = {
            'gross_margin': None,
            'net_profit_margin': None,
            'current_ratio': None,
            'debt_to_equity': None
        }
        
        try:
            income = financials.get('income_stmt', {})
            balance = financials.get('balance_sheet', {})
            
            if not income and not balance:
                logger.warning("No financial data available for ratio computation")
                return ratios
            
            # Get most recent period (first column in yfinance data)
            if income:
                income_df = pd.DataFrame(income)
                if not income_df.empty and len(income_df.columns) > 0:
                    latest = income_df.iloc[:, 0]
                    
                    # Gross Margin
                    if 'Gross Profit' in latest.index and 'Total Revenue' in latest.index:
                        gross_profit = safe_float_conversion(latest.get('Gross Profit'))
                        revenue = safe_float_conversion(latest.get('Total Revenue'))
                        if gross_profit and revenue and revenue != 0:
                            ratios['gross_margin'] = (gross_profit / revenue) * 100
                    
                    # Net Profit Margin
                    if 'Net Income' in latest.index and 'Total Revenue' in latest.index:
                        net_income = safe_float_conversion(latest.get('Net Income'))
                        revenue = safe_float_conversion(latest.get('Total Revenue'))
                        if net_income and revenue and revenue != 0:
                            ratios['net_profit_margin'] = (net_income / revenue) * 100
            
            if balance:
                balance_df = pd.DataFrame(balance)
                if not balance_df.empty and len(balance_df.columns) > 0:
                    latest = balance_df.iloc[:, 0]
                    
                    # Current Ratio
                    if 'Current Assets' in latest.index and 'Current Liabilities' in latest.index:
                        current_assets = safe_float_conversion(latest.get('Current Assets'))
                        current_liabilities = safe_float_conversion(latest.get('Current Liabilities'))
                        if current_assets and current_liabilities and current_liabilities != 0:
                            ratios['current_ratio'] = current_assets / current_liabilities
                    
                    # Debt to Equity
                    if 'Total Debt' in latest.index and 'Stockholders Equity' in latest.index:
                        total_debt = safe_float_conversion(latest.get('Total Debt'))
                        equity = safe_float_conversion(latest.get('Stockholders Equity'))
                        if total_debt and equity and equity != 0:
                            ratios['debt_to_equity'] = total_debt / equity
            
            logger.info(f"Computed ratios: {ratios}")
            return ratios
            
        except Exception as e:
            logger.error(f"Error computing ratios: {e}")
            return ratios

# ---------- Visualization ----------
class ChartGenerator:
    """Handles chart generation."""
    
    @staticmethod
    def create_price_chart(df: pd.DataFrame, ticker: str, output_path: str) -> None:
        """
        Create a price chart with moving averages.
        
        Args:
            df: DataFrame with price data
            ticker: Stock ticker
            output_path: Path to save chart
        """
        try:
            plt.figure(figsize=(12, 6))
            
            # Plot price and moving averages
            plt.plot(df['Date'], df['close'], label='Close Price', linewidth=2)
            if '30d_ma' in df.columns:
                plt.plot(df['Date'], df['30d_ma'], label='30-day MA', alpha=0.7)
            if '50d_ma' in df.columns:
                plt.plot(df['Date'], df['50d_ma'], label='50-day MA', alpha=0.7)
            
            plt.title(f'{ticker} Price History', fontsize=14, fontweight='bold')
            plt.xlabel('Date', fontsize=12)
            plt.ylabel('Price ($)', fontsize=12)
            plt.legend()
            plt.grid(True, alpha=0.3)
            plt.tight_layout()
            
            # Ensure directory exists
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            plt.savefig(output_path, dpi=150, bbox_inches='tight')
            plt.close()
            
            logger.info(f"Created chart: {output_path}")
            
        except Exception as e:
            logger.error(f"Failed to create chart for {ticker}: {e}")
            raise

# ---------- LLM Interface ----------
class LLMInterface:
    """Handles all LLM interactions with retry logic."""
    
    def __init__(self, config: Config):
        self.config = config
        self.llm = ChatOpenAI(
            api_key=config.openai_api_key,
            base_url=config.openai_base_url,
            model=config.model_name,
            temperature=0.0,
            timeout=config.request_timeout
        )
        
        # Define prompts
        self.input_parser_prompt = PromptTemplate(
            input_variables=["user_request"],
            template="""You are an assistant that extracts structured parameters for a financial report.
Input: {user_request}

Return ONLY valid JSON with these keys:
- tickers: list of ticker symbols (strings, uppercase)
- period: one of [1d, 5d, 1mo, 3mo, 6mo, 1y, 2y, 5y, 10y, ytd, max]
- metrics: list of metrics requested (e.g., ["revenue growth", "P/E ratio", "margins"])
- output_format: "markdown" or "pdf" or "docx"
- extras: optional dict for additional parameters

Example:
{{"tickers": ["AAPL", "MSFT"], "period": "1y", "metrics": ["revenue growth", "P/E ratio"], "output_format": "markdown"}}

Return ONLY the JSON object, no other text."""
        )
        
        self.writer_prompt = PromptTemplate(
            input_variables=["context_json", "key_findings", "charts"],
            template="""You are a senior financial analyst writing a concise report.

Context: {context_json}
Key Findings: {key_findings}
Charts Available: {charts}

Write a clear, professional Markdown report (2-4 pages) with:

1. **Executive Summary** (2-4 sentences highlighting main findings)

2. **Key Metrics Table** (format as markdown table with ticker, latest price, return, volatility, ratios)

3. **Analysis** (one paragraph per ticker discussing trends and performance)

4. **Visualizations** (reference chart files by name, e.g., "See AAPL_close.png")

5. **Explainability** - "How I Reached These Conclusions":
   - List specific numerical values that led to conclusions
   - Reference exact metrics (e.g., "AAPL's 30-day MA of $XX.XX crossed above...")
   
6. **Risks & Opportunities** (3-5 bullet points)

Be specific with numbers. Avoid vague statements. Return ONLY the markdown content."""
        )
        
        self.review_prompt = PromptTemplate(
            input_variables=["draft_markdown", "data_summary"],
            template="""Review this financial report draft for accuracy and clarity.

Data Summary: {data_summary}

Draft:
{draft_markdown}

Check for:
1. Numeric consistency with data
2. Ambiguous or unsupported claims
3. Missing context or explanations

Return valid JSON with:
- "issues": list of strings (empty if none found)
- "revised": the corrected markdown (or original if no changes needed)

Return ONLY the JSON object."""
        )
        
        # Create chains
        self.parser_chain = self.input_parser_prompt | self.llm
        self.writer_chain = self.writer_prompt | self.llm
        self.review_chain = self.review_prompt | self.llm
    
    def parse_request(self, user_request: str) -> ParsedRequest:
        """
        Parse natural language request into structured format.
        
        Args:
            user_request: Natural language request
            
        Returns:
            ParsedRequest object
            
        Raises:
            ParsingError: If parsing fails
        """
        logger.info("Parsing user request")
        
        for attempt in range(self.config.max_retries):
            try:
                response = self.parser_chain.invoke({"user_request": user_request})
                parsed_text = response.content.strip()
                
                # Extract JSON if wrapped in markdown code blocks
                if parsed_text.startswith("```"):
                    lines = parsed_text.split("\n")
                    parsed_text = "\n".join(lines[1:-1])
                
                parsed_json = json.loads(parsed_text)
                
                # Normalize output format string
                output_fmt_str = parsed_json.get("output_format", "markdown").lower().strip()
                
                # Map to enum with validation
                format_mapping = {
                    "markdown": OutputFormat.MARKDOWN,
                    "md": OutputFormat.MARKDOWN,
                    "pdf": OutputFormat.PDF,
                    "docx": OutputFormat.DOCX,
                    "doc": OutputFormat.DOCX,
                    "word": OutputFormat.DOCX
                }
                
                output_format = format_mapping.get(output_fmt_str, OutputFormat.MARKDOWN)
                logger.info(f"Parsed output format: '{output_fmt_str}' -> {output_format}")
                
                # Validate and create ParsedRequest
                parsed = ParsedRequest(
                    tickers=[t.upper().strip() for t in parsed_json.get("tickers", [])],
                    period=normalize_period(parsed_json.get("period")),
                    metrics=parsed_json.get("metrics", ["revenue growth", "margins"]),
                    output_format=output_format,
                    extras=parsed_json.get("extras")
                )
                
                parsed.validate()
                logger.info(f"Successfully parsed request: {parsed.tickers}")
                return parsed
                
            except json.JSONDecodeError as e:
                logger.warning(f"JSON decode failed (attempt {attempt + 1}): {e}")
                if attempt == self.config.max_retries - 1:
                    raise ParsingError(f"Failed to parse JSON response: {str(e)}")
            
            except Exception as e:
                logger.error(f"Parsing error (attempt {attempt + 1}): {e}")
                if attempt == self.config.max_retries - 1:
                    raise ParsingError(f"Failed to parse request: {str(e)}")
        
        raise ParsingError("Max retries exceeded")
    
    def generate_report(self, context: Dict, findings: Dict, charts: List[str]) -> str:
        """Generate report markdown."""
        logger.info("Generating report")
        
        try:
            response = self.writer_chain.invoke({
                "context_json": json.dumps(context, indent=2),
                "key_findings": json.dumps(findings, indent=2),
                "charts": json.dumps(charts)
            })
            
            return response.content.strip()
            
        except Exception as e:
            logger.error(f"Failed to generate report: {e}")
            raise
    
    def review_report(self, draft: str, data_summary: Dict) -> tuple[List[str], str]:
        """Review and revise report."""
        logger.info("Reviewing report")
        
        try:
            response = self.review_chain.invoke({
                "draft_markdown": draft,
                "data_summary": json.dumps(data_summary, indent=2)
            })
            
            review_text = response.content.strip()
            
            # Extract JSON from markdown code blocks if present
            if review_text.startswith("```"):
                lines = review_text.split("\n")
                review_text = "\n".join(lines[1:-1])
            
            review_json = json.loads(review_text)
            issues = review_json.get("issues", [])
            revised = review_json.get("revised", draft)
            
            if issues:
                logger.warning(f"Review found {len(issues)} issues")
            else:
                logger.info("Review found no issues")
            
            return issues, revised
            
        except Exception as e:
            logger.error(f"Review failed: {e}")
            return [], draft

# ---------- PDF Conversion ----------
class PDFConverter:
    """Converts markdown reports to PDF format."""
    
    @staticmethod
    def markdown_to_pdf(markdown_content: str, output_path: str, chart_paths: List[str] = None) -> None:
        """
        Convert markdown content to PDF with embedded charts.
        
        Args:
            markdown_content: Markdown text
            output_path: Path to save PDF
            chart_paths: Optional list of chart image paths to embed
            
        Raises:
            RuntimeError: If PDF libraries not available
        """
        if not PDF_AVAILABLE:
            raise RuntimeError(
                "PDF conversion requires additional libraries. "
                "Install with: pip install markdown reportlab"
            )
        
        try:
            logger.info(f"Converting markdown to PDF: {output_path}")
            
            from reportlab.platypus import Image as RLImage
            from reportlab.lib.enums import TA_CENTER
            
            # Create PDF
            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Styles
            styles = getSampleStyleSheet()
            story = []
            
            # Track embedded charts
            charts_added = set()
            chart_paths_set = set(chart_paths) if chart_paths else set()
            
            # Split by lines
            lines = markdown_content.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 0.2*inch))
                    continue
                
                # Clean line of special characters that cause encoding issues
                # Replace bullet points and special chars
                line_clean = line.replace('•', '*').replace('–', '-').replace('—', '-')
                line_clean = line_clean.replace('"', '"').replace('"', '"')
                line_clean = line_clean.replace(''', "'").replace(''', "'")
                
                # Headers
                if line_clean.startswith('# '):
                    p = Paragraph(line_clean[2:], styles['Heading1'])
                    story.append(p)
                    story.append(Spacer(1, 0.2*inch))
                elif line_clean.startswith('## '):
                    p = Paragraph(line_clean[3:], styles['Heading2'])
                    story.append(p)
                    story.append(Spacer(1, 0.15*inch))
                elif line_clean.startswith('### '):
                    p = Paragraph(line_clean[4:], styles['Heading3'])
                    story.append(p)
                    story.append(Spacer(1, 0.1*inch))
                # Lists - use asterisk instead of bullet
                elif line_clean.startswith('- ') or line_clean.startswith('* '):
                    p = Paragraph(f"* {line_clean[2:]}", styles['BodyText'])
                    story.append(p)
                # Check for chart references and embed images
                elif chart_paths_set:
                    chart_embedded = False
                    for chart_path in chart_paths_set:
                        chart_name = Path(chart_path).name
                        if chart_name in line and chart_path not in charts_added:
                            # Add text reference
                            p = Paragraph(line_clean, styles['BodyText'])
                            story.append(p)
                            story.append(Spacer(1, 0.1*inch))
                            
                            # Embed chart image
                            if Path(chart_path).exists():
                                img = RLImage(chart_path, width=6*inch, height=3.5*inch)
                                story.append(img)
                                story.append(Spacer(1, 0.2*inch))
                                charts_added.add(chart_path)
                                chart_embedded = True
                                break
                    
                    if not chart_embedded and not line_clean.startswith('|') and not line_clean.startswith('```'):
                        p = Paragraph(line_clean, styles['BodyText'])
                        story.append(p)
                # Regular text (skip table markers and code blocks)
                else:
                    if not line_clean.startswith('|') and not line_clean.startswith('```'):
                        p = Paragraph(line_clean, styles['BodyText'])
                        story.append(p)
            
            # Build PDF
            doc.build(story)
            logger.info(f"PDF created successfully: {output_path}")
            
        except Exception as e:
            logger.error(f"Failed to convert to PDF: {e}")
            raise

class DOCXConverter:
    """Converts markdown reports to DOCX format."""
    
    @staticmethod
    def markdown_to_docx(markdown_content: str, output_path: str, chart_paths: List[str] = None) -> None:
        """
        Convert markdown content to DOCX with embedded charts.
        
        Args:
            markdown_content: Markdown text
            output_path: Path to save DOCX
            chart_paths: Optional list of chart image paths to embed
            
        Raises:
            RuntimeError: If DOCX libraries not available
        """
        if not DOCX_AVAILABLE:
            raise RuntimeError(
                "DOCX conversion requires python-docx library. "
                "Install with: pip install python-docx"
            )
        
        try:
            logger.info(f"Converting markdown to DOCX: {output_path}")
            
            # Create document
            doc = Document()
            
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Track if we've added charts
            charts_added = set()
            chart_paths_set = set(chart_paths) if chart_paths else set()
            
            # Parse markdown line by line
            lines = markdown_content.split('\n')
            in_code_block = False
            in_table = False
            table_data = []
            
            i = 0
            while i < len(lines):
                line = lines[i].rstrip()
                
                # Handle code blocks
                if line.startswith('```'):
                    in_code_block = not in_code_block
                    i += 1
                    continue
                
                if in_code_block:
                    # Add code as monospace
                    p = doc.add_paragraph(line)
                    p.style = 'Normal'
                    run = p.runs[0] if p.runs else p.add_run()
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    i += 1
                    continue
                
                # Handle tables
                if line.startswith('|') and '|' in line[1:]:
                    if not in_table:
                        in_table = True
                        table_data = []
                    
                    # Parse table row
                    cells = [cell.strip() for cell in line.split('|')[1:-1]]
                    table_data.append(cells)
                    i += 1
                    
                    # Check if next line is separator or end of table
                    if i < len(lines):
                        next_line = lines[i].strip()
                        if next_line.startswith('|---') or not next_line.startswith('|'):
                            if next_line.startswith('|---'):
                                i += 1  # Skip separator
                                continue
                            else:
                                # End of table, create it
                                if table_data:
                                    DOCXConverter._create_table(doc, table_data)
                                    table_data = []
                                in_table = False
                    else:
                        # End of document
                        if table_data:
                            DOCXConverter._create_table(doc, table_data)
                        in_table = False
                    continue
                
                # Empty line
                if not line:
                    if not in_table:
                        doc.add_paragraph()
                    i += 1
                    continue
                
                # Headers
                if line.startswith('# '):
                    p = doc.add_heading(line[2:], level=1)
                    i += 1
                    continue
                elif line.startswith('## '):
                    p = doc.add_heading(line[3:], level=2)
                    i += 1
                    continue
                elif line.startswith('### '):
                    p = doc.add_heading(line[4:], level=3)
                    i += 1
                    continue
                elif line.startswith('#### '):
                    p = doc.add_heading(line[5:], level=4)
                    i += 1
                    continue
                
                # Bullet lists
                if line.startswith('- ') or line.startswith('* '):
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                    i += 1
                    continue
                
                # Numbered lists
                if len(line) > 2 and line[0].isdigit() and line[1] == '.':
                    p = doc.add_paragraph(line[3:], style='List Number')
                    i += 1
                    continue
                
                # Check for chart references and embed images
                if chart_paths_set:
                    for chart_path in chart_paths_set:
                        chart_name = Path(chart_path).name
                        if chart_name in line and chart_path not in charts_added:
                            # Add the paragraph with reference
                            p = doc.add_paragraph(line)
                            DOCXConverter._apply_inline_formatting(p)
                            
                            # Embed the chart image
                            if Path(chart_path).exists():
                                doc.add_picture(chart_path, width=Inches(6))
                                last_paragraph = doc.paragraphs[-1]
                                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                charts_added.add(chart_path)
                            
                            i += 1
                            continue
                
                # Regular paragraph with inline formatting
                p = doc.add_paragraph()
                DOCXConverter._add_formatted_text(p, line)
                i += 1
            
            # Save document
            doc.save(output_path)
            logger.info(f"DOCX created successfully: {output_path}")
            
        except Exception as e:
            logger.error(f"Failed to convert to DOCX: {e}")
            raise
    
    @staticmethod
    def _create_table(doc, table_data: List[List[str]]) -> None:
        """Create a table in the document."""
        if not table_data:
            return
        
        # Create table
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Light Grid Accent 1'
        
        # Populate table
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = cell_text
                
                # Make header row bold
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
    
    @staticmethod
    def _add_formatted_text(paragraph, text: str) -> None:
        """Add text with inline formatting (bold, italic, code)."""
        # Simple parsing for **bold**, *italic*, `code`
        import re
        
        # Pattern to match **bold**, *italic*, or `code`
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
        parts = re.split(pattern, text)
        
        for part in parts:
            if not part:
                continue
            
            if part.startswith('**') and part.endswith('**'):
                # Bold
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                # Italic
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                # Code
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            else:
                # Regular text
                paragraph.add_run(part)
    
    @staticmethod
    def _apply_inline_formatting(paragraph) -> None:
        """Apply inline formatting to existing paragraph text."""
        # This is a simplified version - for production, merge with _add_formatted_text
        pass

# ---------- Main Orchestrator ----------
class FinancialReportOrchestrator:
    """Main orchestrator for financial report generation."""
    
    def __init__(self, config: Config):
        self.config = config
        self.fetcher = DataFetcher(timeout=config.request_timeout)
        self.analyzer = FinancialAnalyzer()
        self.chart_gen = ChartGenerator()
        self.llm = LLMInterface(config)
    
    def analyze_ticker(self, ticker: str, period: str, output_dir: Path) -> TickerAnalysis:
        """
        Analyze a single ticker.
        
        Args:
            ticker: Stock ticker
            period: Time period
            output_dir: Output directory
            
        Returns:
            TickerAnalysis object
        """
        logger.info(f"Analyzing {ticker}")
        
        try:
            # Fetch price data
            df_prices = self.fetcher.fetch_price_history(ticker, period)
            
            # Compute metrics
            df_analyzed = self.analyzer.compute_metrics(df_prices)
            
            # Save CSV
            csv_path = output_dir / f"{ticker}_prices.csv"
            df_analyzed.to_csv(csv_path, index=False)
            
            # Create chart
            chart_path = output_dir / f"{ticker}_close.png"
            self.chart_gen.create_price_chart(df_analyzed, ticker, str(chart_path))
            
            # Fetch and analyze financials
            financials = self.fetcher.fetch_financials(ticker)
            ratios = self.analyzer.compute_ratios(financials)
            
            # Extract key metrics
            latest_close = float(df_analyzed['close'].iloc[-1])
            avg_return = float(df_analyzed['daily_return'].mean())
            volatility = float(df_analyzed['volatility'].iloc[-1]) if not pd.isna(df_analyzed['volatility'].iloc[-1]) else 0.0
            
            # Sample data for review
            sample_data = df_analyzed.tail(3).to_dict(orient="records")
            
            return TickerAnalysis(
                ticker=ticker,
                csv_path=str(csv_path),
                chart_path=str(chart_path),
                latest_close=latest_close,
                avg_daily_return=avg_return,
                volatility=volatility,
                ratios=ratios,
                sample_data=sample_data
            )
            
        except Exception as e:
            logger.error(f"Failed to analyze {ticker}: {e}")
            return TickerAnalysis(
                ticker=ticker,
                csv_path="",
                chart_path="",
                latest_close=0.0,
                avg_daily_return=0.0,
                volatility=0.0,
                ratios={},
                sample_data=[],
                error=str(e)
            )
    
    def run(self, user_request: str, output_dir: str = "./out") -> ReportMetadata:
        """
        Main execution method.
        
        Args:
            user_request: Natural language request
            output_dir: Output directory path
            
        Returns:
            ReportMetadata object
        """
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        
        logger.info("=" * 60)
        logger.info("Starting Financial Report Generation")
        logger.info("=" * 60)
        
        # Step 1: Parse request
        try:
            parsed = self.llm.parse_request(user_request)
        except ParsingError as e:
            logger.error(f"Failed to parse request: {e}")
            raise
        
        # Step 2: Analyze each ticker
        analyses: Dict[str, TickerAnalysis] = {}
        chart_files = []
        
        for ticker in parsed.tickers:
            analysis = self.analyze_ticker(ticker, parsed.period, output_path)
            analyses[ticker] = analysis
            
            if not analysis.error and analysis.chart_path:
                chart_files.append(analysis.chart_path)
        
        # Check if any analyses succeeded
        successful_analyses = [a for a in analyses.values() if not a.error]
        if not successful_analyses:
            raise AnalysisError("All ticker analyses failed")
        
        # Step 3: Prepare context and findings
        context = {
            "tickers": parsed.tickers,
            "period": parsed.period,
            "metrics": parsed.metrics
        }
        
        findings = {
            "analyses": {
                ticker: {
                    "latest_close": analysis.latest_close,
                    "avg_daily_return": analysis.avg_daily_return,
                    "volatility": analysis.volatility,
                    "ratios": analysis.ratios
                }
                for ticker, analysis in analyses.items()
                if not analysis.error
            }
        }
        
        # Step 4: Generate report
        draft_markdown = self.llm.generate_report(context, findings, chart_files)
        
        # Step 5: Review report
        data_summary = {
            ticker: {
                "csv": analysis.csv_path,
                "chart": analysis.chart_path,
                "latest_close": analysis.latest_close,
                "ratios": analysis.ratios,
                "error": analysis.error
            }
            for ticker, analysis in analyses.items()
        }
        
        issues, final_markdown = self.llm.review_report(draft_markdown, data_summary)
        
        # Step 6: Save final report
        timestamp = datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')
        
        # Save markdown version (always create this first)
        md_filename = f"financial_report_{timestamp}.md"
        md_path = output_path / md_filename
        
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(final_markdown)
        
        logger.info(f"Markdown report saved to: {md_path}")
        
        # Convert to PDF if requested
        final_report_path = str(md_path)
        
        if parsed.output_format == OutputFormat.PDF:
            pdf_filename = f"financial_report_{timestamp}.pdf"
            pdf_path = output_path / pdf_filename
            
            try:
                converter = PDFConverter()
                converter.markdown_to_pdf(final_markdown, str(pdf_path), chart_files)
                final_report_path = str(pdf_path)
                logger.info(f"PDF report saved to: {pdf_path}")
            except RuntimeError as e:
                logger.warning(f"PDF conversion failed: {e}. Returning markdown instead.")
            except Exception as e:
                logger.error(f"Unexpected error during PDF conversion: {e}. Returning markdown.")
        
        elif parsed.output_format == OutputFormat.DOCX:
            docx_filename = f"financial_report_{timestamp}.docx"
            docx_path = output_path / docx_filename
            
            try:
                converter = DOCXConverter()
                converter.markdown_to_docx(final_markdown, str(docx_path), chart_files)
                final_report_path = str(docx_path)
                logger.info(f"DOCX report saved to: {docx_path}")
            except RuntimeError as e:
                logger.warning(f"DOCX conversion failed: {e}. Returning markdown instead.")
            except Exception as e:
                logger.error(f"Unexpected error during DOCX conversion: {e}. Returning markdown.")
        
        logger.info(f"Final report: {final_report_path}")
        logger.info("=" * 60)
        logger.info("Financial Report Generation Complete")
        logger.info("=" * 60)
        
        return ReportMetadata(
            final_markdown_path=final_report_path,
            charts=chart_files,
            analyses=analyses,
            review_issues=issues,
            generated_at=timestamp
        )

# ---------- Main Entry Point ----------
def main():
    """Main entry point for script execution."""
    try:
        # Load configuration
        config = Config.from_env()
        
        # Create orchestrator
        orchestrator = FinancialReportOrchestrator(config)
        
        # Example request
        user_request = """
        Generate a financial report for AAPL and MSFT covering the last year.
        Include revenue growth, profit margins, P/E ratio, and volatility analysis.
        Output format: pdf.
        """
        
        # Run orchestration
        result = orchestrator.run(user_request, output_dir="./report_out")
        
        # Print results
        print("\n" + "=" * 60)
        print("REPORT GENERATION SUMMARY")
        print("=" * 60)
        print(f"Report: {result.final_markdown_path}")
        print(f"Format: {result.final_markdown_path.split('.')[-1].upper()}")
        print(f"Charts: {len(result.charts)} generated")
        print(f"Tickers analyzed: {len(result.analyses)}")
        
        if result.review_issues:
            print(f"\nReview Issues Found: {len(result.review_issues)}")
            for i, issue in enumerate(result.review_issues, 1):
                print(f"  {i}. {issue}")
        else:
            print("\nNo review issues found")
        
        print("\nAnalysis Results:")
        for ticker, analysis in result.analyses.items():
            if analysis.error:
                print(f"  {ticker}: ERROR - {analysis.error}")
            else:
                print(f"  {ticker}: Success")
                print(f"    Latest Close: ${analysis.latest_close:.2f}")
                print(f"    Avg Return: {analysis.avg_daily_return*100:.2f}%")
                print(f"    Volatility: {analysis.volatility*100:.2f}%")
        
        print("=" * 60)
        
    except Exception as e:
        logger.error(f"Fatal error: {e}", exc_info=True)
        raise

if __name__ == "__main__":
    main()
